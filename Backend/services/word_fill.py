import re
import unicodedata
import copy
from pathlib import Path
from typing import Dict, Optional, Callable
from decimal import Decimal
from docx import Document
from docx.table import Table

__all__ = ["fill_docx"]

# =========================
#  Placeholders y formatos
# =========================

PLACEHOLDER_RX = re.compile(
    r"\{\{\s*([^{}|]+?)\s*(?:\|\s*([a-z0-9_]+)\s*)?\}\}"
)

def _strip_accents(s: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFKD", s)
        if not unicodedata.combining(c)
    )

def _norm_key(s: str) -> str:
    s = str(s or "")
    s = _strip_accents(s).lower()
    s = s.replace(" ", "")
    return s

def _fmt_moneda(x) -> str:
    try:
        d = Decimal(str(x or 0))
        return f"${d:,.0f}".replace(",", ".")
    except Exception:
        return str(x)

def _fmt_entero(x) -> str:
    try:
        d = int(Decimal(str(x or 0)))
        return f"{d:,}".replace(",", ".")
    except Exception:
        return str(x)

def _fmt_dec2(x) -> str:
    try:
        d = Decimal(str(x or 0))
        return f"{d:,.2f}".replace(",", ".")
    except Exception:
        return str(x)

_FORMATTERS: Dict[str, Callable[[object], str]] = {
    "moneda": _fmt_moneda,
    "entero": _fmt_entero,
    "dec2": _fmt_dec2,
}

# =========================
#  Lookup de contexto
#  (SIN metas/productos)
# =========================

# Prefijos que NO queremos reemplazar todavía
_META_PROD_PREFIXES = [
    "meta",                # meta, meta_1, meta_indicador, etc.
    "cod_meta",
    "numero_meta",
    "nombre_meta",
    "producto",
    "cod_producto",
    "indicador_producto",
    "cod_indicador_producto",
    "meta_indicador",
]

def _is_meta_prod_key(nk: str) -> bool:
    return any(nk.startswith(p) for p in _META_PROD_PREFIXES)

def _build_lookup(context: Dict[str, object]) -> Dict[str, str]:
    out: Dict[str, str] = {}

    for k, v in (context or {}).items():
        nk = _norm_key(k)              # ej: "cod_meta_1" -> "cod_meta_1"
        out[nk] = "" if v is None else str(v)

        # Alias: nombre_meta_X -> meta_X
        if nk.startswith("nombre_meta_"):
            suf = nk.split("nombre_meta_")[-1]
            out[_norm_key(f"meta_{suf}")] = out[nk]

        # Alias: numero_meta_X -> cod_meta_X
        if nk.startswith("numero_meta_"):
            suf = nk.split("numero_meta_")[-1]
            out[_norm_key(f"cod_meta_{suf}")] = out[nk]

        # Espejos día/dia
        if "dia" in nk and "texto" in nk:
            out[nk.replace("dia", "día")] = out[nk]
        if "día" in k or "dia" in k:
            out[_norm_key(k.replace("día", "dia"))] = out[nk]
            out[_norm_key(k.replace("dia", "día"))] = out[nk]

    return out

# =========================
#  Reemplazo de texto
# =========================

def _replace_text(text: str, lookup: Dict[str, str]) -> str:
    def repl(m):
        raw_key = m.group(1)
        raw_fmt = m.group(2)
        nk = _norm_key(raw_key)

        if nk not in lookup:
            # si no está en lookup, dejamos el placeholder tal cual
            return m.group(0)

        val = lookup[nk]
        if raw_fmt and raw_fmt in _FORMATTERS:
            try:
                num_val = Decimal(
                    val.replace(".", "").replace("$", "").replace(",", ".")
                )
            except Exception:
                num_val = val
            return _FORMATTERS[raw_fmt](num_val)
        return val

    return PLACEHOLDER_RX.sub(repl, text)

def _replace_in_run_level(paragraph, lookup: Dict[str, str]) -> bool:
    joined = "".join(r.text for r in paragraph.runs)
    if not PLACEHOLDER_RX.search(joined):
        return True
    for r in paragraph.runs:
        if PLACEHOLDER_RX.search(r.text or ""):
            r.text = _replace_text(r.text, lookup)
    after = "".join(r.text for r in paragraph.runs)
    return not PLACEHOLDER_RX.search(after)

def _replace_in_paragraph(paragraph, lookup: Dict[str, str]) -> None:
    if _replace_in_run_level(paragraph, lookup):
        return
    full_text = "".join(r.text for r in paragraph.runs)
    new_text = _replace_text(full_text, lookup)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)

def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

# =========================
#  Utilidades de tablas
# =========================

def _renumber_placeholders_in_cell(cell, idx: int):
    META_KEYS = {
        "cod_meta",
        "meta",
        "numero_meta",
        "nombre_meta",
        "producto",
        "cod_producto",
        "indicador_producto",
        "cod_indicador_producto",
        "meta_indicador",
    }
    META_KEYS_NORM = {_norm_key(k) for k in META_KEYS}

    def renumber_text(text: str) -> str:
        def repl(m):
            raw_key = m.group(1)   # "cod_meta_1" o "meta" o "producto_1"
            raw_fmt = m.group(2)   # formato opcional, ej: "moneda"
            key = raw_key.strip()

            # separar posible sufijo numérico
            parts = key.split("_")
            base = key
            if parts[-1].isdigit():
                base = "_".join(parts[:-1])  # "cod_meta_1" -> "cod_meta"

            base_norm = _norm_key(base)

            # solo tocamos las claves de metas/productos
            if base_norm in META_KEYS_NORM:
                new_key = f"{base}_{idx}"
            else:
                new_key = key

            if raw_fmt:
                return "{{" + new_key + "|" + raw_fmt + "}}"
            else:
                return "{{" + new_key + "}}"

        return PLACEHOLDER_RX.sub(repl, text)

    for p in cell.paragraphs:
        full_text = "".join(r.text or "" for r in p.runs)
        new_text = renumber_text(full_text)

        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new_text)


def _clone_row_after(row, table):
    tr = row._tr
    new_tr = copy.deepcopy(tr)
    tr.addnext(new_tr)
    for r in table.rows:
        if r._tr is new_tr:
            return r
    return table.rows[-1]

def _insert_table_after(existing_table):
    src_tbl = existing_table._tbl
    new_tbl_elm = copy.deepcopy(src_tbl)
    src_tbl.addnext(new_tbl_elm)
    return Table(new_tbl_elm, existing_table._parent)

# =========================
#  Expansión de CARTA:
#  Metas y Productos
# =========================

def expand_metas_in_carta(doc: Document, total_metas: int) -> None:
    if total_metas <= 1:
        return

    for tbl in doc.tables:
        meta_rows = []
        for row in tbl.rows:
            text = " || ".join(c.text for c in row.cells)
            if "{{cod_meta_1}}" in text or "{{meta_1}}" in text:
                meta_rows.append(row)

        if not meta_rows:
            continue

        plantilla = meta_rows
        anchor = meta_rows[-1]

        for i in range(2, total_metas + 1):
            for r in plantilla:
                new_r = _clone_row_after(anchor, tbl)
                anchor = new_r
                for cell in new_r.cells:
                    _renumber_placeholders_in_cell(cell, i)

        # solo queremos procesar la primera tabla que tenga este bloque
        break

def expand_productos_in_carta(doc: Document, total_productos: int) -> None:
    if total_productos <= 1:
        return

    tablas_detectadas = []

    for tbl in doc.tables:
        texto_tabla = " ".join(
            c.text for row in tbl.rows for c in row.cells
        )
        if "{{producto_1}}" in texto_tabla:
            tablas_detectadas.append(tbl)

    if not tablas_detectadas:
        return

    base_tbl = tablas_detectadas[0]

    duplicated_tables = [base_tbl]
    anchor = base_tbl
    for _ in range(2, total_productos + 1):
        new_tbl = _insert_table_after(anchor)
        duplicated_tables.append(new_tbl)
        anchor = new_tbl

    for idx, tbl in enumerate(duplicated_tables, start=1):
        for row in tbl.rows:
            for cell in row.cells:
                _renumber_placeholders_in_cell(cell, idx)

# =========================
#  FUNCIÓN PRINCIPAL
# =========================

def fill_docx(base_dir: Path, template_name: str, context: Dict[str, object], output_name: Optional[str] = None, output_dir: Optional[Path] = None,) -> Path:
    template_path = base_dir / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"No se encontró el template: {template_path}")

    doc = Document(str(template_path))

    # --- 1) Expansión de metas/productos SOLO si viene __metas_ctx__
    metas_ctx = context.get("__metas_ctx__") or []
    total_metas = len(metas_ctx)

    if total_metas > 0:
        expand_metas_in_carta(doc, total_metas)
        expand_productos_in_carta(doc, total_metas)  # 1 producto por meta

    # --- 2) Reemplazo del resto de variables (NO metas/productos)
    lookup = _build_lookup(context)

    for p in _iter_all_paragraphs(doc):
        _replace_in_paragraph(p, lookup)

    # --- 3) Guardar
    out_dir = output_dir or base_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / (output_name or f"filled_{template_name}")
    doc.save(str(out_path))
    return out_path
